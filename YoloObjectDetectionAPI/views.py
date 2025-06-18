
import base64
import io
import json
import logging
import os
import tempfile
import traceback
import uuid
import zipfile
from datetime import timedelta
from io import BytesIO

import cv2
import matplotlib.pyplot as plt
import numpy as np
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.platypus import Table, TableStyle
from docx import Document
from docx.shared import Inches

from django.conf import settings
from django.contrib import messages
from django.utils import timezone
from django.utils.html import strip_tags
from django.utils.timezone import localtime

from django.contrib.auth import login, logout
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib.auth.forms import AuthenticationForm
from django.contrib.auth.models import User

from django.core.files import File
from django.core.files.base import ContentFile
from django.core.mail import EmailMessage, send_mail
from django.db.models import Avg, Count
from django.db.models.functions import TruncDate
from django.http import (
    FileResponse, Http404, HttpResponse, HttpResponseServerError, JsonResponse
)
from django.shortcuts import get_object_or_404, render, redirect
from django.template.loader import render_to_string
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST

from rest_framework import status
from rest_framework.response import Response
from rest_framework.views import APIView

from .forms import (
    CameraForm, CustomAuthenticationForm, CustomUserCreationForm, OrderForm
)
from .models import (
    Camera, CameraStream, Detections, ImageUpload, Order
)
from .serializers import ImageUploadSerializer
from .detection_models.yolov8 import YOLOv8Detector


#Пароль для тест аккааунтов Test123!

logger = logging.getLogger(__name__)

DEFECT_SUMMARY = {
    "Отслаивание (деламинация)": {
        "Описание": "Появление дыр, расслоение фигуры во время печати.",
        "Причина": "Неоднородная полимерная масса (плохое хранение/долгое отстаивание).",
        "Что делать": [
            "Проводить тестовую печать и испытания смолы перед серийной печатью.",
            "Хранить смолу в соответствии с рекомендациями производителя.",
            "Перемешивать смолу перед использованием."
        ]
    },
    "Деформация формы/креплений": {
        "Описание": "Модель получает нежелательные изгибы, искажение поддержек.",
        "Причина": "Недостаточная проработка модели или ошибка в количестве поддержек, возможно вторично — из-за отслаивания.",
        "Что делать": [
            "Проверить 3D-модель на целостность и адекватность поддержек перед печатью.",
            "Убедиться в корректных настройках слайсера.",
            "Учитывать возможность вторичных дефектов от других причин (например, отслаивания)."
        ]
    },
    "Неправильно засохшие крепления": {
        "Описание": "Форма модели сохраняется, но теряется текстура поверхности.",
        "Причина": "Неравномерная сушка, деформация после отверждения.",
        "Что делать": [
            "Улучшить контроль УФ-отверждения.",
            "Использовать равномерную постобработку.",
            "Контролировать поддерживающие структуры при проектировании."
        ]
    },
    "Провисания и геометрические искажения": {
        "Описание": "Длинные и тонкие элементы провисают или отламываются.",
        "Причина": "Неполное отверждение, недостаточная УФ-обработка.",
        "Что делать": [
            "Увеличить время УФ-обработки.",
            "Проверять мощность и равномерность источника УФ.",
            "Избегать слишком длинных элементов без поддержек."
        ]
    },
    "Неправильная настройка печатной платформы": {
        "Описание": "Основание модели не прикрепляется, ось наклонена, повреждается экран.",
        "Причина": "Неровность осей платформы, смещение под весом.",
        "Что делать": [
            "Проверить и откалибровать платформу перед каждой печатью.",
            "Следить за равномерным распределением веса модели.",
            "Немедленно остановить печать при обнаружении неровностей."
        ]
    }
}

@csrf_exempt
def register_view(request):
    if request.method == 'POST':
        form = CustomUserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)

            subject = f"Новый заказ от пользователя {user.username}"
            message = f"Здравствуйте!\n\nПользователь {user.username} зарегистрировался."
            try:
                send_result = send_mail(
                    subject,
                    message,
                    settings.DEFAULT_FROM_EMAIL,
                    ['polinastorozh@gmail.com'],
                    message
                )
                print(f"📧 Письмо отправлено? Результат: {send_result}")
            except Exception as e:
                print(f"❌ Ошибка при отправке email: {str(e)}")

            return index(request, register_form=CustomUserCreationForm(), login_form=AuthenticationForm(), registration_success=True)
    else:
        form = CustomUserCreationForm()

    return index(request, register_form=form, login_form=AuthenticationForm())


def login_view(request):
    if request.method == 'POST':
        form = CustomAuthenticationForm(request, data=request.POST)
        if form.is_valid():
            user = form.get_user()
            login(request, user)
            return redirect('yolov8:home')
    else:
        form = CustomAuthenticationForm()

    return index(request, login_form=form, register_form=CustomUserCreationForm())


def logout_view(request):
    logout(request)
    return redirect('yolov8:home')

MAX_FILE_SIZE_MB = 500
ALLOWED_EXTENSIONS = ['.obj', '.fbx', '.stl', '.zip', '.rar']


@login_required
@csrf_exempt
def order_view(request):
    if request.method == 'POST':
        form = OrderForm(request.POST)
        uploaded_file = request.FILES.get('model_file')

        if not uploaded_file:
            return JsonResponse({"success": False, "error": "Файл не прикреплён"}, status=400)

        ext = os.path.splitext(uploaded_file.name)[1].lower()
        if ext not in ['.obj', '.fbx', '.stl']:
            return JsonResponse({"success": False, "error": "Неподдерживаемый формат файла. Разрешены: .obj, .fbx, .stl"}, status=400)

        if uploaded_file.size > 500 * 1024 * 1024:
            return JsonResponse({"success": False, "error": "Файл превышает 500 МБ"}, status=400)

        if form.is_valid():
            order = form.save(commit=False)
            order.user = request.user
            order.model_file = uploaded_file 
            order.save()
            subject = f"Создан новый заказ: {order.order_name}"
            message = f"Пользователь {request.user.username} создал заказ «{order.order_name}» (ID: {order.id})."

            try:
                send_result = send_mail(
                    subject,
                    message,
                    settings.DEFAULT_FROM_EMAIL,
                    ['polinastorozh@gmail.com'],
                )
                print(f"📧 Уведомление отправлено? Результат: {send_result}")
            except Exception as e:
                print(f"❌ Ошибка при отправке email: {str(e)}")

            return JsonResponse({"success": True, "order_name": order.order_name})

        return JsonResponse({"success": False, "errors": form.errors}, status=400)

    return JsonResponse({"success": False, "error": "Только POST-запросы разрешены"}, status=405)


def upload_image(request):
    images = ImageUpload.objects.all()
    context = {'images': images}
    return render(request, 'process_image.html', context)


class UploadImageView(APIView):
    detector = YOLOv8Detector('last.pt')

    def post(self, request, format=None):
        image_data = request.data.get('image_data')
        order_id = request.data.get('order_id')

        if not image_data:
            return JsonResponse({'error': 'No image data provided'}, status=400)

        try:
            image_data = image_data.split(',')[1]
            img_bytes = base64.b64decode(image_data)
            img_array = np.frombuffer(img_bytes, dtype=np.uint8)
            img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

            random_filename = str(uuid.uuid4())
            media_dir = os.path.join('media', 'uploads')
            os.makedirs(media_dir, exist_ok=True)

            original_filename = f'{random_filename}_original.jpg'
            processed_filename = f'{random_filename}_processed.jpg'

            original_img_path = os.path.join(media_dir, original_filename)
            processed_img_path = os.path.join(media_dir, processed_filename)

            cv2.imwrite(original_img_path, img)

            with open(original_img_path, 'rb') as tmp_file:
                django_file = File(tmp_file)
                image_upload = ImageUpload(
                    status=ImageUpload.STATUS_PENDING,
                    order_id=order_id,
                    confidence_threshold=0.5
                )
                image_upload.image_file.save(original_filename, django_file, save=True)

            confidence_threshold = image_upload.confidence_threshold
            detection_result = self.detector.run_detection(
                image_path=original_img_path,
                confidence_threshold=confidence_threshold
            )

            for det in detection_result:
                Detections.objects.create(
                    object_detection=image_upload,
                    label=det['label'],
                    confidence=det['confidence'],
                    x_min=det['x_min'],
                    x_max=det['x_max'],
                    y_min=det['y_min'],
                    y_max=det['y_max'],
                )

            label_translations = {
                'NoRow': 'Нет ряда',
                'Defect': 'Дефект',
                'NoDefect': 'Без дефекта',
                'Row': 'Ряд'
            }

            img_pil = Image.fromarray(cv2.cvtColor(img, cv2.COLOR_BGR2RGB))
            draw = ImageDraw.Draw(img_pil)

            font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"  
            if not os.path.exists(font_path):
                font_path = "arial.ttf"  
            font = ImageFont.truetype(font_path, 24)

            for det in detection_result:
                x_min = int(det['x_min'])
                y_min = int(det['y_min'])
                x_max = int(det['x_max'])
                y_max = int(det['y_max'])
                label = det['label']
                label_ru = label_translations.get(label, label)

                if label in ['NoRow', 'Defect']:
                    color = (255, 0, 0) 
                elif label in ['NoDefect', 'Row']:
                    color = (0, 255, 0) 
                else:
                    color = (0, 255, 255) 
                draw.rectangle([x_min, y_min, x_max, y_max], outline=color, width=3)
                draw.text((x_min, y_min - 25), label_ru, font=font, fill=color)

            img = cv2.cvtColor(np.array(img_pil), cv2.COLOR_RGB2BGR)
            cv2.imwrite(processed_img_path, img)

            with open(processed_img_path, 'rb') as tmp_file:
                django_file = File(tmp_file)
                image_upload.image_file.save(processed_filename, django_file, save=True)
                image_upload.status = ImageUpload.STATUS_COMPLETED
                image_upload.processed_timestamp = timezone.now()
                image_upload.save()

            _, img_encoded = cv2.imencode('.jpg', img)
            img_base64 = base64.b64encode(img_encoded).decode('utf-8')

            serialized_detections = []
            for det in detection_result:
                serialized_detections.append({
                    'label': str(det['label']),
                    'confidence': float(det['confidence']),
                    'x_min': int(det['x_min']),
                    'x_max': int(det['x_max']),
                    'y_min': int(det['y_min']),
                    'y_max': int(det['y_max']),
                })

            return JsonResponse({
                'processed_image_url': f"data:image/jpeg;base64,{img_base64}",
                'detections': serialized_detections,
                'confidence_threshold': confidence_threshold
            })

        except Exception as e:
            traceback_str = traceback.format_exc()
            print(traceback_str)
            return JsonResponse({'error': str(e), 'trace': traceback_str}, status=500)

def index(request, form=None, register_form=None, login_form=None, registration_success=False):
    total_images = ImageUpload.objects.all().count()
    images_processed = ImageUpload.objects.filter(status=ImageUpload.STATUS_COMPLETED).count()
    images_pending = ImageUpload.objects.filter(status=ImageUpload.STATUS_PROCESSING).count()

    defect_count = Detections.objects.count()
    defect_count_by_label = Detections.objects.values('label').annotate(count=Count('label'))

    images_processed_data = [total_images, images_processed, images_pending]
    defects_found_data = [defect['count'] for defect in defect_count_by_label]

    if not form:
        form = OrderForm()

    if not register_form:
        register_form = CustomUserCreationForm()

    if not login_form:
        login_form = AuthenticationForm()

    context = {
        'total_images': total_images,
        'images_processed': images_processed,
        'images_pending': images_pending,
        'user': request.user,
        'images_processed_data': images_processed_data,
        'defects_found_data': defects_found_data,
        'form': form,
        'register_form': register_form,
        'login_form': login_form,
        'registration_success': registration_success

    }
    return render(request, 'home/index.html', context)

def defect_analysis(request):
    try:
        image_statuses = ImageUpload.objects.values('status').annotate(count=Count('status'))
        defect_types = Detections.objects.values('label').annotate(count=Count('label'))
        defect_dates = Detections.objects.annotate(day=TruncDate('detection_date')).values('day').annotate(count=Count('day')).order_by('day')
        average_defects_per_image = Detections.objects.aggregate(average_defects=Avg('object_detection'))

        statuses = {status['status']: status['count'] for status in image_statuses}
        labels = [defect['label'] for defect in defect_types]
        defect_counts = [defect['count'] for defect in defect_types]
        dates = [date['day'] for date in defect_dates]
        counts = [date['count'] for date in defect_dates]
        average_defects = average_defects_per_image['average_defects'] or 0

        return render(request, 'defect_analysis.html', {
            'statuses': statuses,
            'labels': labels,
            'defect_counts': defect_counts,
            'dates': dates,
            'counts': counts,
            'average_defects': average_defects
        })

    except Exception as e:
        logger.error(f"Ошибка при анализе дефектов: {e}", exc_info=True)
        return HttpResponseServerError("Произошла ошибка при обработке данных.")

def is_admin(user):
    return user.is_staff

@login_required
@user_passes_test(is_admin)
def admin_panel(request):
    if request.user.is_staff:
        orders = Order.objects.all()
        return render(request, 'admin_panel.html', {'orders': orders})
    else:
        return redirect('home')

@login_required
@user_passes_test(is_admin)
def update_order_status(request, order_id, status):
    try:
        order = Order.objects.get(id=order_id)
        order.status = status
        order.save()
        return JsonResponse({'success': True, 'status': order.get_status_display()})
    except Order.DoesNotExist:
        return JsonResponse({'success': False, 'message': 'Order not found'})

@login_required
def client_orders(request):
    orders = Order.objects.filter(user=request.user)
    return render(request, 'client_orders.html', {'orders': orders})

def camera_page(request):
    orders = Order.objects.all().order_by('-created_at')
    return render(request, 'camera_page.html', {'orders': orders})



@login_required
def download_model_file(request, order_id):
    try:
        order = Order.objects.get(pk=order_id)
        if order.model_file:
            file = order.model_file.open('rb')
            filename = os.path.basename(order.model_file.name)
            return FileResponse(file, as_attachment=True, filename=filename)
        else:
            raise Http404("Файл не найден")
    except Order.DoesNotExist:
        raise Http404("Заказ не найден")

def get_order_details(request, order_id):
    try:
        order = Order.objects.get(id=order_id)
        return JsonResponse({
            'order_name': order.order_name,
            'description': order.description,
            'material': order.material,
            'status': order.get_status_display(),
            'created_at': order.created_at.strftime('%Y-%m-%d %H:%M'),
        })
    except Order.DoesNotExist:
        return JsonResponse({'error': 'Заказ не найден'}, status=404)


def is_admin(user):
    return user.is_staff

@login_required
@user_passes_test(is_admin)
def admin_panel(request):
    if request.user.is_staff:
        orders = Order.objects.all()
        camera, created = Camera.objects.get_or_create(name="Default Camera", defaults={'camera_id': 0})

        context = {
            'orders': orders,
            'websocket_url': "ws/camera/"
        }
        return render(request, 'admin_panel.html', context)
    else:
        return redirect('home')

@require_POST
@login_required
@user_passes_test(is_admin)
def delete_order(request, order_id):
    try:
        order = Order.objects.get(id=order_id)
        order.delete()
        return JsonResponse({'success': True})
    except Order.DoesNotExist:
        return JsonResponse({'success': False, 'message': 'Заказ не найден'})


@csrf_exempt
def update_order_status(request):
    try:
        data = json.loads(request.body)
        order_id = data.get('order_id')
        new_status = data.get('status')

        order = get_object_or_404(Order, id=order_id)

        old_status = order.status
        if old_status == new_status:
            return JsonResponse({'success': False, 'message': 'Статус не изменился.'})

        order.status = new_status
        order.save()

        client_email = order.user.email
        print(f" Заказ ID {order_id}: статус изменён с '{old_status}' на '{new_status}'")
        print(f" Email клиента: {client_email}")

        if client_email:
            subject = f"Обновление статуса вашего заказа: {order.order_name}"
            message = f"Здравствуйте!\n\nСтатус вашего заказа «{order.order_name}» обновлён на: {order.get_status_display()}."

            try:
                send_result = send_mail(
                    subject,
                    message,
                    settings.DEFAULT_FROM_EMAIL,
                    [client_email],
                )
                print(f" Письмо отправлено? Результат: {send_result}")
            except Exception as e:
                print(f" Ошибка при отправке email: {str(e)}")

        return JsonResponse({'success': True, 'status': order.get_status_display()})

    except Exception as e:
        print(f" Ошибка при обновлении статуса: {str(e)}")
        return JsonResponse({'success': False, 'message': str(e)})

def defect_analysis(request):
    clients = User.objects.filter(orders__isnull=False).distinct()
    return render(request, 'defect_analysis.html', {'clients': clients})

def get_orders(request, user_id):
    orders = Order.objects.filter(user_id=user_id).exclude(status='not_started').values('id', 'order_name', 'status')
    return JsonResponse(list(orders), safe=False)

def defect_analysis_client(request, user_id):
    client = get_object_or_404(User, id=user_id)
    orders = Order.objects.filter(client=client)
    return render(request, 'client_defect_analysis.html', {'client': client, 'orders': orders})

def get_orders_client(request, user_id):
    orders = Order.objects.filter(client__id=user_id).values('id', 'order_name', 'status')
    return JsonResponse(list(orders), safe=False)

def get_images_client(request, order_id):
    images = Image.objects.filter(order_id=order_id).values('image_file', 'status')
    return JsonResponse(list(images), safe=False)

@login_required
def get_order_images(request, order_id):
    order = get_object_or_404(Order, id=order_id, user=request.user)
    images = ImageUpload.objects.filter(order=order)

    data = [
        {
            'image_url': image.image_file.url,
            'status': image.get_status_display(),
            'upload_timestamp': image.upload_timestamp.strftime('%Y-%m-%d %H:%M:%S') if image.upload_timestamp else '',
            'processed_timestamp': image.processed_timestamp.strftime('%Y-%m-%d %H:%M:%S') if image.processed_timestamp else ''
        }
        for image in images
    ]
    return JsonResponse({'images': data})

def get_orders_client(request, user_id):
    orders = Order.objects.filter(client__id=user_id).values('id', 'order_name', 'status')
    return JsonResponse(list(orders), safe=False)

def get_images_client(request, order_id):
    images = Image.objects.filter(order_id=order_id).values('image_file', 'status')
    return JsonResponse(list(images), safe=False)

def get_images(request, order_id):
    images = ImageUpload.objects.filter(order_id=order_id).values('id', 'image_file', 'status')

    for image in images:
        image['image_file'] = f"/media/{image['image_file']}"

    return JsonResponse(list(images), safe=False)

def get_order_summary(request, order_id):
    try:
        order = Order.objects.get(id=order_id)
    except Order.DoesNotExist:
        return JsonResponse({'error': 'Заказ не найден'}, status=404)

    images = ImageUpload.objects.filter(order=order)
    detections = Detections.objects.filter(object_detection__in=images)

    duration = order.updated_at - order.created_at
    duration_seconds = int(duration.total_seconds())
    duration_days = duration_seconds // 86400
    duration_hours = (duration_seconds % 86400) // 3600
    duration_minutes = (duration_seconds % 3600) // 60

    days_str = f"{duration_days} {'день' if duration_days == 1 else 'дня' if 2 <= duration_days <= 4 else 'дней'}" if duration_days else ''
    hours_str = f"{duration_hours} {'час' if duration_hours == 1 else 'часа' if 2 <= duration_hours <= 4 else 'часов'}" if duration_hours else ''
    minutes_str = f"{duration_minutes} {'минута' if duration_minutes == 1 else 'минуты' if 2 <= duration_minutes <= 4 else 'минут'}" if duration_minutes else ''

    duration_str = ', '.join(filter(None, [days_str, hours_str, minutes_str]))

    order_summary_info = [
        f"Название заказа: {order.order_name}",
        f"Заказчик: {order.user.username}",
        f"Материал: {order.material}",
        f"Статус: {order.get_status_display()}",
        f"Дата создания: {localtime(order.created_at).strftime('%Y-%m-%d %H:%M')}",
        f"Последнее обновление: {localtime(order.updated_at).strftime('%Y-%m-%d %H:%M')}",
        f"Продолжительность выполнения: {duration_str}",
        f"Количество изображений: {images.count()}",
        f"Количество детекций: {detections.count()}",
    ]

    if not detections.exists():
        return JsonResponse({
            'summary': order_summary_info + ['По данному заказу не найдено детекций.'],
            'chartData': {},
            'detectionInfo': []
        })

    summary_data = detections.values('label').annotate(
        count=Count('id'),
        avg_confidence=Avg('confidence')
    )

    explanation_map = {
        'NoRow': 'Нет ряда',
        'Defect': 'Дефекты',
        'Row': 'Ровный ряд',
        'NoDefect': 'Нет дефектов',
    }

    detailed_summary = []
    labels = []
    counts = []
    confidences = []
    detection_info = []
    scatter_data = []
    
    trend_data_area = []
    trend_data_confidence = []
    trend_data_avg_confidence = []

    for entry in summary_data:
        label = entry['label']
        count = entry['count']
        avg_conf = round(entry['avg_confidence'] * 100, 2)
        explanation = explanation_map.get(label, 'Обнаружен дефект.')
        detailed_summary.append(f"{label} — {count} шт. (уверенность: {avg_conf}%) → {explanation}")
        labels.append(label)
        counts.append(count)
        confidences.append(avg_conf)

    for detection in detections:
        area = (detection.x_max - detection.x_min) * (detection.y_max - detection.y_min)
        detection_info.append({
            'label': detection.label,
            'confidence': detection.confidence,
            'bounding_box': {
                'x_min': detection.x_min,
                'y_min': detection.y_min,
                'x_max': detection.x_max,
                'y_max': detection.y_max,
                'area': area
            },
            'detection_date': detection.detection_date.isoformat()
        })

        scatter_data.append({
            'x': detection.detection_date.isoformat(),
            'y': area
        })

        trend_data_area.append({
            'x': detection.detection_date.isoformat(),
            'y': area
        })
        trend_data_confidence.append({
            'x': detection.detection_date.isoformat(),
            'y': detection.confidence * 100
        })

        avg_confidence = detections.filter(label=detection.label).aggregate(avg_confidence=Avg('confidence'))['avg_confidence'] * 100
        trend_data_avg_confidence.append({
            'x': detection.detection_date.isoformat(),
            'y': avg_confidence
        })

    return JsonResponse({
        'summary': order_summary_info + [''] + detailed_summary,
        'chartData': {
            'labels': labels,
            'counts': counts,
            'confidences': confidences,
            'scatter': scatter_data,
            'area_trend': trend_data_area,
            'confidence_trend': trend_data_confidence,
            'avg_confidence_trend': trend_data_avg_confidence
        },
        'detectionInfo': detection_info
    })


def get_chart_data(order):
    return {
        'labels': ['Defect', 'NoDefect', 'Row', 'NoRow'],
        'counts': [15, 8, 22, 3]
    }

def calculate_area(bbox):
    """Функция для вычисления площади по координатам bounding box"""
    return (bbox['x_max'] - bbox['x_min']) * (bbox['y_max'] - bbox['y_min'])

def generate_defect_graph(order):
    chart_data = get_chart_data(order)
    fig, ax = plt.subplots(figsize=(6, 3))
    ax.bar(chart_data['labels'], chart_data['counts'], color='#ff6384')
    ax.set_title('Количество дефектов по категориям')
    ax.set_xlabel('Тип дефекта')
    ax.set_ylabel('Количество')
    img_buffer = BytesIO()
    fig.savefig(img_buffer, format='png')
    img_buffer.seek(0)
    
    return img_buffer

def download_summary(request, order_id):
    order = get_object_or_404(Order, id=order_id)

    images = ImageUpload.objects.filter(order=order)
    detections = Detections.objects.filter(object_detection__in=images)

    document = Document()
    document.add_heading(f"Сводка по заказу: {order.order_name}", 0)

    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Параметр'
    hdr_cells[1].text = 'Значение'
    
    order_summary_info = [
        ('Заказчик', order.user.username),
        ('Материал', order.material),
        ('Статус', order.get_status_display()),
        ('Дата создания', localtime(order.created_at).strftime('%Y-%m-%d %H:%M')),
        ('Последнее обновление', localtime(order.updated_at).strftime('%Y-%m-%d %H:%M')),
        ('Количество изображений', str(images.count())),
        ('Количество детекций', str(detections.count())),
    ]
    
    for param, value in order_summary_info:
        row_cells = table.add_row().cells
        row_cells[0].text = param
        row_cells[1].text = value
    img_buffer = generate_defect_graph(order)
    document.add_picture(img_buffer, width=Inches(5))
    
    if detections.exists():
        document.add_heading('Подробности по детекциям:', level=1)
        detection_table = document.add_table(rows=1, cols=4)
        detection_table.style = 'Table Grid'
        detection_hdr_cells = detection_table.rows[0].cells
        detection_hdr_cells[0].text = 'Тип'
        detection_hdr_cells[1].text = 'Уверенность (%)'
        detection_hdr_cells[2].text = 'Площадь (px²)'
        detection_hdr_cells[3].text = 'Границы'
        
        for detection in detections:
            area = (detection.x_max - detection.x_min) * (detection.y_max - detection.y_min)
            
            row_cells = detection_table.add_row().cells
            row_cells[0].text = detection.label
            row_cells[1].text = str(detection.confidence)
            row_cells[2].text = str(area)
            row_cells[3].text = f"X: {detection.x_min} - {detection.x_max}, Y: {detection.y_min} - {detection.y_max}"
    
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    return FileResponse(buffer, as_attachment=True, filename=f"order_{order.id}_summary.docx")

def download_images(request, order_id):
    order = get_object_or_404(Order, id=order_id)
    images = ImageUpload.objects.filter(order=order)

    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "w") as zip_file:
        for image in images:
            filename = image.image_file.name.split("/")[-1]
            zip_file.writestr(filename, image.image_file.read())

    zip_buffer.seek(0)
    return FileResponse(zip_buffer, as_attachment=True, filename='images.zip')

def generate_summary_pdf(filepath, summary_lines):
    c = canvas.Canvas(filepath)
    text = c.beginText(50, 800)
    text.setFont("Helvetica", 12)
    for line in summary_lines:
        text.textLine(line)
    c.drawText(text)
    c.save()

@login_required
def order_page(request):
    orders = Order.objects.filter(user=request.user)
    return render(request, 'order_page.html', {'orders': orders})


def client_defect_analysis(request, client_id):
    client = User.objects.get(id=client_id)
    orders = Order.objects.filter(user_id=client_id).exclude(status='not_started') 
    status_dict = {
        'not_started': 'Не начато',
        'in_progress': 'В процессе',
        'completed': 'Завершено'
    }

    orders_with_translated_status = [
        {
            'order': order,
            'translated_status': status_dict.get(order.status, order.status)
        }
        for order in orders
    ]

    return render(request, 'client_defect_analysis.html', {
        'client': client,
        'orders_with_translated_status': orders_with_translated_status
    })

def camera_list(request):
    cameras = Camera.objects.all()
    form = CameraForm()
    return render(request, 'camera_list.html', {'cameras': cameras, 'form': form})

def add_camera(request):
    if request.method == 'POST':
        form = CameraForm(request.POST)
        if form.is_valid():
            camera = form.save()
            return JsonResponse({'success': True, 'id': camera.id, 'name': camera.name})
    return JsonResponse({'success': False})

def delete_camera(request, pk):
    camera = get_object_or_404(Camera, pk=pk)
    camera.delete()
    return JsonResponse({'success': True})

def available_cameras(request):
    cameras = Camera.objects.filter(is_active=False)
    data = [{'id': cam.id, 'name': cam.name} for cam in cameras]
    return JsonResponse(data, safe=False)

@csrf_exempt
def toggle_camera_status(request):
    if request.method == 'POST':
        body = json.loads(request.body)
        cam_id = body.get('camera_id')
        is_active = body.get('is_active')
        try:
            camera = Camera.objects.get(id=cam_id)
            camera.is_active = is_active
            camera.save()
            return JsonResponse({'success': True})
        except Camera.DoesNotExist:
            return JsonResponse({'success': False, 'message': 'Camera not found'})
    return JsonResponse({'success': False, 'message': 'Invalid request'})

@csrf_exempt
def get_or_create_stream(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        order_id = data.get('order_id')
        camera_id = data.get('camera_id')

        stream, created = CameraStream.objects.get_or_create(
            order_id=order_id,
            camera_id=camera_id,
            is_active=True,
            defaults={}
        )
        return JsonResponse({'stream_url': f'/yolov8/stream/{stream.stream_id}/'})

def watch_stream(request, stream_id):
    return render(request, 'stream_view.html', {'stream_id': stream_id})
